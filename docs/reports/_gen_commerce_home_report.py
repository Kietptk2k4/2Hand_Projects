"""Generate DOCX report for commerce-home-hybrid-ltr (from OpenSpec artifacts only)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


OUT = Path(
    r"d:\Projects\2Hand_Projects\docs\reports\Bao-cao-de-xuat-Commerce-Home-hybrid-LTR.docx"
)


def set_run_font(run, *, bold=False, size=11):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, size=14 if level == 1 else 12)
    return h


def add_para(doc, text, *, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=10)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "BÁO CÁO CHỨC NĂNG ĐỀ XUẤT COMMERCE HOME\n"
        "HYBRID LTR (RULES + ENTITY CF + CROSS-DOMAIN AR + LIGHTGBM + DIVERSITY)"
    )
    set_run_font(run, bold=True, size=16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Nguồn: OpenSpec change commerce-home-hybrid-ltr\n"
        "(proposal.md, design.md, specs/*)\n"
        "Dự án 2Hands"
    )
    set_run_font(run, size=11)

    add_heading(doc, "1. Mục đích và phạm vi", 1)
    add_para(
        doc,
        "Báo cáo này mô tả chức năng đề xuất sản phẩm trên Commerce Home theo đúng "
        "OpenSpec change commerce-home-hybrid-ltr. Commerce Home trước đây chủ yếu "
        "là duyệt catalog (sort/filter), chưa có đường đề xuất cá nhân hóa do "
        "commerce-service sở hữu. Mục tiêu là xây hybrid recommend: Rules + Entity CF "
        "+ Cross-domain AR + LightGBM ranking + Diversity, trả Top 50.",
    )
    add_para(doc, "Trong phạm vi (theo proposal):", bold=True)
    add_bullets(
        doc,
        [
            "Commerce Home hybrid recommend thuộc commerce-service.",
            "Auth-only: bắt buộc JWT; không có guest/anonymous Home recommend.",
            "Retrieval A ∪ B ∪ E ∪ C (không semantic fill v1).",
            "Entity-based CF (không product_id→product_id làm đồ thị chính).",
            "Cross-domain AR tag→category + export social interest.",
            "Impression/engage async; click v1 = product-detail attributed from=home.",
            "LTR binary; HOME_FEATURE_ORDER 15 chiều; PopularityNormalizer từ artifact.",
            "Train data modes SEED_ONLY | HYBRID | REAL_ONLY qua Admin system_configs.",
            "Model riêng commerce_home_ranker trong Commerce model_artifacts.",
        ],
    )
    add_para(doc, "Ngoài phạm vi (theo proposal):", bold=True)
    add_bullets(
        doc,
        [
            "Post-feed retrain orchestrator/drift; ALS/BPR; Neo4j.",
            "Stage-1 rule ranker cho Home production Top-K.",
            "Viewport impression beacons; guest Home rail.",
            "Semantic ANN retrieval (D0 — deferred).",
            "Post-style shared-DB seed-db cho Home; dedicated Docker sim DB v1.",
        ],
    )

    add_heading(doc, "2. Định nghĩa chiến lược Hybrid", 1)
    add_para(
        doc,
        "Theo architecture note và proposal: Hybrid Commerce Home = Rules (Source A) "
        "+ Entity CF (Source E) + Cross-domain AR (Source C) + LightGBM LTR + Diversity. "
        "Semantic retrieval (Source D) bị hoãn. LightGBM chỉ chấm điểm pool đã có, "
        "không invent candidate ngoài pool.",
    )

    add_heading(doc, "3. Hợp đồng API phục vụ (Online)", 1)
    add_para(doc, "3.1 Endpoint", bold=True)
    add_bullets(
        doc,
        [
            "GET /commerce/api/v1/home/recommendations",
            "Owner: commerce-service; không gọi Social recommend-feed trên hot path.",
            "Trả tối đa 50 sản phẩm ACTIVE còn hàng.",
            "Payload success gồm request_id và ranking_mode ∈ {LIGHTGBM, DEGRADED}.",
            "Mỗi item tối thiểu: id, title, price, thumbnail, shop (minimal), rating (summary hoặc null).",
            "Giá listing dùng effective_price (Commerce domain); không invent taxonomy giá Home riêng.",
        ],
    )
    add_para(doc, "3.2 Hành vi biên", bold=True)
    add_table(
        doc,
        ["Tình huống", "Hành vi bắt buộc"],
        [
            ["Unauthenticated", "Từ chối (ví dụ 401); không trả popular-only guest Top K"],
            ["Feature flag tắt (COMMERCE_HOME_RECOMMEND_ENABLED)", "HTTP 404 + mã feature-disabled ổn định"],
            ["Pool rỗng sau filter", "HTTP 200 với items = []"],
            ["ONNX/normalizer thiếu", "ranking_mode = DEGRADED; key 0.7*popularity + 0.3*recency"],
        ],
    )

    add_heading(doc, "4. Pipeline online (điểm phục vụ)", 1)
    add_para(
        doc,
        "Thứ tự online (specs commerce-home-hybrid-recommend + candidate-retrieval):",
    )
    add_bullets(
        doc,
        [
            "1) Xác thực JWT; kiểm tra feature flag.",
            "2) Build UserInterestProfile tại as_of = now.",
            "3) Retrieval A/B/E/C → CandidateProduct; hard filter inventory + own-shop + vacation; pool ≤ 500.",
            "4) Feature Builder → vector 15 chiều (HOME_FEATURE_ORDER).",
            "5) LightGBM/ONNX score (hoặc degraded key); sort score DESC, tie created_at DESC rồi product_id ASC.",
            "6) Diversity greedy hard-cap (config); backfill nếu thiếu K.",
            "7) Trả Top K (mặc định 50); async log impression + provenance.",
        ],
    )

    add_heading(doc, "5. UserInterestProfile", 1)
    add_bullets(
        doc,
        [
            "Commerce facets (category, brand, shop) và social facets (hashtag, keyword) tách biệt — không gộp social tag vào category score.",
            "Commerce: cửa sổ 180 ngày; COMPLETED ×1.0; cart ×0.6; decay 2^(-Δ/30d); max-norm; giữ top 20 categories, 20 brands, 10 shops.",
            "Cart MVP: ưu tiên bảng cart-add event; nếu không có thì cart_items.created_at trong cửa sổ (loại soft-deleted).",
            "Social facets: đọc user_social_interest_export (max-norm theo tag_type); không gọi Social HTTP trên hot path.",
            "Price percentiles p25/p50/p75 từ COMPLETED effective_price trong 180d; nếu < 3 mẫu → missing (price_affinity default).",
            "Chỉ dùng event nghiêm ngặt trước as_of.",
        ],
    )

    add_heading(doc, "6. Candidate retrieval (A/B/E/C)", 1)
    add_para(
        doc,
        "Soft caps: A=100, B=150, E=150, C=150. Dedupe ≤ 500 CandidateProduct; OR-merge sources. "
        "Không semantic/ANN (D). Pool có thể < 500 — 500 là upper bound, không phải fill target.",
    )
    add_para(doc, "6.1 CandidateProduct", bold=True)
    add_bullets(
        doc,
        [
            "Snapshot sản phẩm đủ cho Feature Builder (category, brand, shop, price, created_at, rating…).",
            "Tập nguồn + personalScore / cfScore / arScore (nullable).",
            "cfScore, arScore đã ∈ [0,1] khi có; Feature Builder chỉ clip phòng thủ.",
            "Cấm shared retrievalScore xuyên nguồn.",
        ],
    )
    add_para(doc, "6.2 Source A — Popular/New/Rating (không cá nhân hóa)", bold=True)
    add_bullets(
        doc,
        [
            "Cap 100; slices soft: New 40 (created_at DESC), Popular 40 (completed order-items 90d), Rating 20 (rating_avg với rating_count ≥ 3).",
        ],
    )
    add_para(doc, "6.3 Source B — Personal", bold=True)
    add_bullets(
        doc,
        [
            "Map facet category/brand/shop → sản phẩm ACTIVE in-stock.",
            "personalScore = max(category_score, brand_score, shop_score); cap 150.",
            "Không nhân recency tại retrieval; empty facets → B rỗng.",
        ],
    )
    add_para(doc, "6.4 Source E — Entity CF", bold=True)
    add_bullets(
        doc,
        [
            "Top-N neighbor entities (config); sản phẩm theo completed_order_items DESC rồi created_at DESC.",
            "Soft per-neighbor product cap mặc định 20.",
            "cfScore = raw / max(raw trong request, ε); raw = edge_score × seed_strength.",
        ],
    )
    add_para(doc, "6.5 Source C — Cross-domain AR", bold=True)
    add_bullets(
        doc,
        [
            "Map social tags qua AR confidence ≥ threshold (mặc định 0.05) → category.",
            "Sản phẩm theo category; per-category soft cap mặc định 20.",
            "arScore = clip(confidence × tag_score_norm, 0, 1); lấy max khi nhiều rule.",
        ],
    )
    add_para(doc, "6.6 Hard filters", bold=True)
    add_bullets(
        doc,
        [
            "ACTIVE + còn hàng theo quy tắc sellable Commerce.",
            "Loại sản phẩm shop của chính caller.",
            "Loại sản phẩm shop đang vacation.",
        ],
    )

    add_heading(doc, "7. Entity CF (offline + online)", 1)
    add_bullets(
        doc,
        [
            "Không dùng product_id→product_id làm CF chính (stock=1 second-hand).",
            "Entity v1: LEAF_CATEGORY, BRAND (shop deferred).",
            "Cửa sổ offline 180 ngày trước as_of.",
            "COMPLETED unordered pairs ×1.0; cart co-occur trong 24h ×0.6; không self-pair.",
            "score = log1p(Σ weights); top-M neighbors/entity, mặc định M=50.",
            "Bảng entity_cooccur: (entity_type, entity_id, neighbor_type, neighbor_id, score, updated_at).",
            "HYBRID merge cạnh: max(real_score, seed_score) — không cộng.",
        ],
    )

    add_heading(doc, "8. Social → Commerce AR bridge", 1)
    add_bullets(
        doc,
        [
            "Interest weights: search×4, save×3, comment×2, like×1; cửa sổ 90 ngày; decay export 2^(-Δ/14d).",
            "Bảng user_social_interest_export: user_id, tag_type (HASHTAG|KEYWORD), tag, score, window_days, computed_at, as_of.",
            "Cadence mặc định daily; prefer full refresh per user.",
            "AR: interest_tag → commerce_category only; Apriori (default) hoặc FP-Growth.",
            "Basket 90d: social tags ∪ leaf categories từ COMPLETED.",
            "min_support=0.01, min_confidence=0.05 (config).",
            "Bảng social_tag_category_ar: tag, tag_type, category_id, support, confidence, updated_at.",
            "Online không gọi Social /feed/for-you.",
        ],
    )

    add_heading(doc, "9. Learning-to-Rank (HOME_FEATURE_ORDER)", 1)
    add_para(
        doc,
        "Thứ tự cố định 15 chiều (Python train ≡ Java serve); mismatch → fail closed / DEGRADED:",
    )
    add_bullets(
        doc,
        [
            "1 recency_score — nửa đời 7 ngày: 2^(-Δ/(7·86400)); thiếu created_at → 0",
            "2 popularity_score — raw = COUNT completed order_items (completed_at < as_of); z=log1p(raw); PopularityNormalizer(z_lo,z_hi) từ artifact; cấm min-max trong pool request",
            "3 rating_score — rating_count < 3 → 0.5; else clip(rating_avg/5,0,1)",
            "4 category_match — profile.category_scores[category_id] hoặc 0",
            "5 brand_match — null brand → 0; else brand_scores",
            "6 shop_match — shop_scores",
            "7 price_affinity — theo p25/p75/IQR; missing/IQR≤0 → 0.5; trong [p25,p75] → 1.0",
            "8 cross_domain_score — clip(arScore)",
            "9 cf_score — clip(cfScore)",
            "10 semantic_similarity — 0 khi D disabled",
            "11–15 is_popular, is_personal, is_cf, is_cross_domain, is_semantic (0/1; is_semantic=0 khi D off)",
        ],
    )
    add_para(doc, "Nhãn train:", bold=True)
    add_bullets(
        doc,
        [
            "Binary: click ∪ add-to-cart ∪ purchase trong cửa sổ nhãn (mặc định 24h).",
            "Attribution: nearest prior impression (shown_at ≤ t và còn trong window).",
            "Sample-weight theo loại action: deferred (không bắt buộc v1).",
        ],
    )
    add_para(doc, "Dataset / evaluate:", bold=True)
    add_bullets(
        doc,
        [
            "Một row = một impression; as_of = shown_at; event nghiêm trước as_of.",
            "Tái dựng CandidateProduct từ provenance đã log (sources, personal/cf/ar scores) — không replay retrieval.",
            "Split thời gian theo shown_at: 80/10/10; không shuffle ngẫu nhiên.",
            "Fit PopularityNormalizer trên train only (ưu tiên p1–p99 của log1p(raw)).",
            "Objective: LightGBM binary.",
            "Evaluate: AUC toàn cục + Precision@10 = trung bình precision theo nhóm request_id.",
            "Baseline evaluate = 0.7*popularity_score + 0.3*recency_score (khớp DEGRADED online).",
        ],
    )

    add_heading(doc, "10. Diversity", 1)
    add_bullets(
        doc,
        [
            "Sau scoring: greedy hard-cap theo leaf category (và optional brand/shop) — giá trị cap từ config.",
            "Backfill theo score nếu dưới K.",
            "Có thể tắt diversity → Top K theo score thuần.",
            "Không thay đổi kích thước pool trước score.",
        ],
    )

    add_heading(doc, "11. Impression, click, ModelLoader", 1)
    add_para(doc, "11.1 Impression (async, không block HTTP)", bold=True)
    add_bullets(
        doc,
        [
            "Một row / sản phẩm trả về (chỉ Top K, không full pool).",
            "Fields: user_id, product_id, shown_at, rank_position (sau diversity), request_id, ranking_mode, sources, personal_score, cf_score, ar_score.",
            "sources = JSONB array ∈ {POPULAR, PERSONAL, CF, CROSS_DOMAIN}.",
            "Không suppress-recently-shown; cho phép re-impress.",
            "Không bắt buộc lưu full vector 15 chiều trên impression.",
            "Index: (user_id, shown_at), (user_id, product_id, shown_at), (request_id).",
        ],
    )
    add_para(doc, "11.2 Click", bold=True)
    add_bullets(
        doc,
        [
            "Product-detail với attribution from=home → async CLICK engage.",
            "Thiếu request_id nhưng vẫn attributed từ Home vẫn ghi click.",
            "Detail không attributed → không ghi Home CLICK.",
        ],
    )
    add_para(doc, "11.3 HomeModelLoader", bold=True)
    add_bullets(
        doc,
        [
            "Load ONNX commerce_home_ranker từ MODEL_ROOT + basename portable.",
            "Load kèm feature_order + PopularityNormalizer.",
            "Startup + reload schedule (config); soft-reject là việc offline activate — không ép online activate bản fail gate.",
        ],
    )

    add_heading(doc, "12. Train data mode & Home sim", 1)
    add_table(
        doc,
        ["config_key", "value_type", "Default"],
        [
            ["commerce.home.ltr.train_data_mode", "STRING", "SEED_ONLY"],
            ["commerce.home.ltr.seed_row_weight", "DECIMAL", "0.5"],
            ["commerce.home.ltr.real_only_min_impressions", "INTEGER", "5000"],
        ],
    )
    add_bullets(
        doc,
        [
            "Enum bắt buộc: SEED_ONLY | HYBRID | REAL_ONLY; invalid → fail closed.",
            "Đọc Admin GET /admin/api/v1/system-configs với exact configKey (SYSTEM_CONFIG_VIEW).",
            "Env override chỉ khi RECSYS_HOME_CONFIG_FROM_ENV=1 (CLI smoke).",
            "SEED_ONLY: chỉ file/RAM sim; HYBRID: union real+seed (SEED weight mặc định 0.5); REAL_ONLY: chỉ real, fail nếu impression < min.",
            "Home sim: in-memory catalog từ persona niches; ghi files; cấm INSERT fake users/posts/orders vào Auth/Social/Commerce app DBs.",
            "load-artifact: upsert entity_cooccur, user_social_interest_export, social_tag_category_ar — không pollute event tables.",
            "HYBRID CF edge merge = max(real, seed).",
        ],
    )

    add_heading(doc, "13. Retrain orchestration (offline)", 1)
    add_para(
        doc,
        "Thứ tự bắt buộc (commerce-home-ltr): resolve train_data_mode → (optional Home sim) → "
        "entity CF → social interest export → AR mine → load-artifact (khi cần) → "
        "build-dataset → split → train → evaluate → export-activate.",
    )
    add_bullets(
        doc,
        [
            "Trigger: operator/CLI (cron optional later).",
            "export-activate chỉ sau evaluate.",
            "Gate: lightgbm.auc ≥ baseline.auc AND lightgbm.precision_at_10 ≥ baseline.precision_at_10; null → fail closed.",
            "Soft-reject: không activate; giữ active Home model cũ.",
            "Model name: commerce_home_ranker (tách biệt feed_ranker).",
            "Commerce model_artifacts: id, model_name, version, format, artifact_path, metrics, is_active, trained_at; unique (model_name, version); tối đa một active / model_name.",
            "recsys-offline không được gọi trên online recommend hot path.",
        ],
    )

    add_heading(doc, "14. Tác động hệ thống (Impact)", 1)
    add_table(
        doc,
        ["Thành phần", "Thay đổi theo proposal"],
        [
            ["commerce-service", "Use case recommend, retrieval, feature builder, HomeModelLoader, diversity, API, impression/click"],
            ["recsys-offline", "Home sim files, mode-aware CF/AR/export/LTR, load-artifact, đọc Admin mode"],
            ["admin-service", "Seed/maintain system_configs train mode (+ seed_row_weight); không sở hữu training rows"],
            ["social-service", "Input read-only cho REAL/HYBRID; không sở hữu Home rank; không Home sim write Social DB"],
            ["Frontend/mobile", "Rail Đề xuất → Top 50 API; Admin system-configs sửa train mode; Model Registry phân biệt feed_ranker vs commerce_home_ranker"],
        ],
    )

    add_heading(doc, "15. Capabilities OpenSpec liên quan", 1)
    add_bullets(
        doc,
        [
            "commerce-home-hybrid-recommend — online pipeline + impression/click + degraded",
            "commerce-home-candidate-retrieval — A/B/E/C CandidateProduct",
            "commerce-entity-cf — entity co-occur offline/online",
            "commerce-social-ar-bridge — AR + social export consumer",
            "commerce-home-ltr — formulas, dataset provenance, P@10 by request_id, gate, commerce_home_ranker",
            "commerce-home-train-data-mode — SEED_ONLY/HYBRID/REAL_ONLY + sim + load-artifact",
            "recsys-offline-ops (modified) — mở rộng jobs Home; vẫn không hot-path online",
        ],
    )

    add_heading(doc, "16. Kết luận", 1)
    add_para(
        doc,
        "Chức năng đề xuất Commerce Home theo commerce-home-hybrid-ltr là một đường recommend "
        "do commerce-service sở hữu, auth-only, hybrid retrieval ổn định với hàng second-hand "
        "(entity CF + cross-domain AR + rules), xếp hạng bằng LightGBM trên feature order khóa "
        "15 chiều, có diversity, logging provenance cho train lại, và tách biệt model "
        "commerce_home_ranker khỏi Social feed_ranker. Semantic retrieval bị hoãn; rule-based "
        "chỉ là baseline evaluate / degraded serve — không phải production Top-K ranker.",
    )

    footer = doc.add_paragraph()
    run = footer.add_run(
        "\n— Hết báo cáo. Nội dung bám proposal.md, design.md và các specs trong "
        "openspec/changes/commerce-home-hybrid-ltr/."
    )
    set_run_font(run, size=10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
