"""Generate DOCX report for Social For You feed recommendation (from implemented code)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


OUT = Path(
    r"d:\Projects\2Hand_Projects\docs\reports\Bao-cao-de-xuat-Social-For-You-feed.docx"
)


def set_run_font(run, *, bold=False, size=11):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, size=14 if level == 1 else 12)
    return h


def add_para(doc, text, *, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=10)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "BÁO CÁO CHỨC NĂNG ĐỀ XUẤT BÀI VIẾT\n"
        "SOCIAL FEED “FOR YOU” (ĐỀ XUẤT)\n"
        "CANDIDATE RECALL + 6-DIM FEATURES + LIGHTGBM ONNX / RULE FALLBACK"
    )
    set_run_font(run, bold=True, size=16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Nguồn: mã nguồn đã triển khai (social-service, recsys-offline, frontend)\n"
        "Không mô tả Commerce Home / commerce_home_ranker\n"
        "Dự án 2Hands"
    )
    set_run_font(run, size=11)

    add_heading(doc, "1. Mục đích và phạm vi", 1)
    add_para(
        doc,
        "Báo cáo mô tả chức năng đề xuất bài viết trên tab Social “Đề xuất” (For You) "
        "theo đúng những gì đã implement trong codebase. Owner online là social-service; "
        "train/export offline nằm ở Services/recsys-offline. Đây là đường ranking bài viết "
        "(post feed), khác hoàn toàn Commerce Home hybrid LTR (15 chiều, commerce_home_ranker).",
    )
    add_para(doc, "Trong phạm vi đã implement:", bold=True)
    add_bullets(
        doc,
        [
            "API GET /api/v1/social/feed/for-you — RecommendPostsUseCase.",
            "Candidate recall: followee ưu tiên + public fill; cửa sổ 7/30/90 ngày; loại bài đã xem.",
            "Feature vector 6 chiều (PostFeatureBuilder) khớp FEATURE_ORDER offline.",
            "Ranking LightGBM ONNX model feed_ranker; fallback RuleBasedRankingModel khi thiếu model.",
            "Ghi impression async (post_impression_log) + user_seen_posts.",
            "Offline: clean → build-dataset → split → train → evaluate → export-activate.",
            "Admin Model Registry đọc artifacts/status cho feed_ranker (read-only).",
            "Web FE tab “Đề xuất” gọi đúng /feed/for-you.",
        ],
    )
    add_para(doc, "Ngoài phạm vi / chưa làm (theo code hiện tại):", bold=True)
    add_bullets(
        doc,
        [
            "Diversity / MMR re-rank trên Social For You (chỉ có ở Commerce Home).",
            "Affinity mua hàng real-time online: EmptyUserProductAffinityClient → cross_domain ≈ 0 khi serving.",
            "Bảng engage-event riêng cho feed; label train lấy từ like/save/comment trong 24h sau impression.",
            "LambdaRank / listwise; objective offline là binary.",
            "Admin UI activate/rollback model; chỉ xem registry.",
            "Mobile: tab “Đề xuất” vẫn gọi /feed/global (chronological), chưa wire /for-you.",
            "Không nhầm với GET /feed/global (ViewGlobalFeedUseCase — mới nhất, không ranking).",
        ],
    )

    add_heading(doc, "2. Phân biệt các feed Social", 1)
    add_table(
        doc,
        ["Endpoint", "Use case", "Hành vi"],
        [
            [
                "GET /api/v1/social/feed/for-you",
                "RecommendPostsUseCase",
                "Candidate pool → features → LightGBM/rule → phân trang → log impression",
            ],
            [
                "GET /api/v1/social/feed/global",
                "ViewGlobalFeedUseCase",
                "Chronological PUBLIC; không ranking/model/impression recommend",
            ],
            [
                "GET /api/v1/social/feed/following",
                "ViewFollowingFeedUseCase",
                "Feed theo follow; không dùng feed_ranker",
            ],
        ],
    )
    add_para(
        doc,
        "Response shape của For You tái sử dụng ViewGlobalFeedResult / ViewGlobalFeedResponse "
        "(cùng mapper ViewGlobalFeedHttpMapper) nhưng luồng nghiệp vụ khác Global.",
    )

    add_heading(doc, "3. Hợp đồng API phục vụ (Online)", 1)
    add_para(doc, "3.1 Endpoint và auth", bold=True)
    add_bullets(
        doc,
        [
            "Method/URL: GET /api/v1/social/feed/for-you",
            "Controller: FeedController (@GetMapping(\"/for-you\"))",
            "Auth: JWT bắt buộc; userId null → UNAUTHORIZED",
            "Query: page ≥ 0 (mặc định 0); size ∈ [1, 50] (mặc định 20)",
            "Message thành công: \"Lay recommend feed thanh cong.\"",
        ],
    )
    add_para(doc, "3.2 Owner và model", bold=True)
    add_bullets(
        doc,
        [
            "Owner: social-service",
            "Model name: feed_ranker (social.recommendation.model-name / SOCIAL_RECOMMENDATION_MODEL_NAME)",
            "Bảng registry: Social Postgres model_artifacts",
            "MODEL_ROOT: social.recommendation.model-root ← SOCIAL_RECOMMENDATION_MODEL_ROOT "
            "(Compose thường mount recsys-offline/data/artifacts → /models/recsys)",
        ],
    )

    add_heading(doc, "4. Pipeline online (điểm phục vụ)", 1)
    add_para(
        doc,
        "Thứ tự trong RecommendPostsUseCase (đã implement):",
    )
    add_bullets(
        doc,
        [
            "1) Xác thực JWT.",
            "2) CandidatePoolService.getCandidates(userId, maxSize=500).",
            "3) PostFeatureBuilder.buildFeatureVectors → vector 6 chiều / bài.",
            "4) Nếu social.recommendation.ranking.model=lightgbm (mặc định) và ModelLoader có OrtSession "
            "→ LightGBMRankingModel; ngược lại RuleBasedRankingModel.",
            "5) Sắp xếp score giảm dần; phân trang offset = page * size.",
            "6) Hydrate post + like state; map response.",
            "7) PostImpressionLogger ghi async impression + upsert user_seen_posts cho các item trên trang.",
        ],
    )

    add_heading(doc, "5. Candidate recall", 1)
    add_para(
        doc,
        "Impl: CandidatePoolServiceImpl. Nguồn Mongo posts; seen set từ Postgres user_seen_posts "
        "(lỗi query seen → coi như chưa xem).",
    )
    add_para(doc, "Bộ lọc bài:", bold=True)
    add_bullets(
        doc,
        [
            "status = ACTIVE; visibility = PUBLIC",
            "created_at ≥ now − windowDays",
            "Moderation: thiếu / null / NONE",
            "Loại post_id đã có trong user_seen_posts",
        ],
    )
    add_para(doc, "Nguồn gộp (dedupe):", bold=True)
    add_bullets(
        doc,
        [
            "Followee priority: bài của followee đã accept, limit 300, mới nhất trước.",
            "Public fill: cùng cửa sổ, limit maxSize (500), đổ đầy đến maxSize.",
        ],
    )
    add_para(doc, "Fallback cửa sổ thời gian:", bold=True)
    add_bullets(
        doc,
        [
            "Config: social.recommendation.recall.window-days mặc định 7,30,90 "
            "(SOCIAL_RECOMMENDATION_RECALL_WINDOW_DAYS).",
            "Thử lần lượt; dừng khi poolSize ≥ min-pool-size (mặc định 20, "
            "SOCIAL_RECOMMENDATION_RECALL_MIN_POOL_SIZE).",
            "Half-life feature recency vẫn 7 ngày; cửa sổ recall có thể rộng hơn độc lập.",
        ],
    )

    add_heading(doc, "6. Feature vector (6 chiều)", 1)
    add_para(
        doc,
        "Online: PostFeatureBuilder / PostFeatureVector. Offline: FEATURE_ORDER trong pipelines/train.py. "
        "Thứ tự phải khớp ONNX input float[batch][6]. feature_version = 1.",
    )
    add_table(
        doc,
        ["#", "Tên feature", "Công thức / ghi chú (đã implement)"],
        [
            ["0", "recency_score", "2^(-Δ / 7 ngày)"],
            [
                "1",
                "engagement_score",
                "min-max trên batch của log(1+likes) + 2·log(1+comments)",
            ],
            [
                "2",
                "hashtag_match_score",
                "min-max; trọng số search 1.0 / saved 0.8 / liked 0.4",
            ],
            [
                "3",
                "author_affinity_score",
                "min-max; follow +1 + liked×0.5 + saved×0.6",
            ],
            [
                "4",
                "mutual_follow_score",
                "1 nếu đang follow author; không thì Jaccard tập followee",
            ],
            [
                "5",
                "cross_domain_product_score",
                "0.6·cat_overlap + 0.4·shop_overlap (CrossDomainProductScore)",
            ],
        ],
    )
    add_para(
        doc,
        "Lịch sử online dùng: recent likes 50, saves page 50, search keywords 20. "
        "Client affinity online hiện là EmptyUserProductAffinityClient → "
        "cross_domain_product_score thường ≈ 0 khi serving. Offline vẫn có thể tính "
        "từ export purchase profile (CSV) khi build dataset.",
    )

    add_heading(doc, "7. Ranking: LightGBM ONNX và fallback rule", 1)
    add_para(doc, "7.1 LightGBM / ONNX", bold=True)
    add_bullets(
        doc,
        [
            "Class: LightGBMRankingModel; session từ ModelLoader (ONNX Runtime).",
            "Bật khi ranking.model=lightgbm và OrtSession != null.",
            "ModelLoader: @PostConstruct + @Scheduled cron mặc định 0 0 * * * * "
            "(SOCIAL_RECOMMENDATION_RELOAD_CRON); resolve active feed_ranker dưới MODEL_ROOT "
            "(basename, ví dụ feed_ranker_v3.onnx); fallback SOCIAL_RECOMMENDATION_MODEL_PATH "
            "(mặc định models/recommend_post_model.onnx).",
            "forceReload() có trên ModelLoader; không có HTTP admin force-reload.",
        ],
    )
    add_para(doc, "7.2 Rule-based degraded (khóa Phase-1, tổng = 1.0)", bold=True)
    add_para(
        doc,
        "Class: RuleBasedRankingModel — phải khớp baseline offline evaluate:",
    )
    add_table(
        doc,
        ["Feature", "Trọng số"],
        [
            ["recency", "0.12"],
            ["engagement", "0.28"],
            ["hashtag", "0.22"],
            ["author_affinity", "0.13"],
            ["mutual_follow", "0.13"],
            ["cross_domain", "0.12"],
        ],
    )
    add_para(
        doc,
        "Khi rule-based: impression ghi model_version / model_name = null. "
        "Khi LightGBM: ghi version/name từ artifact đang load. "
        "Không có diversity re-rank trên Social For You.",
    )

    add_heading(doc, "8. Impression và seen", 1)
    add_bullets(
        doc,
        [
            "Bảng post_impression_log: user_id, post_id, shown_at, rank_position, "
            "model_version, model_name, request_id.",
            "Bảng user_seen_posts: (user_id, post_id) — loại khỏi recall sau này.",
            "PostImpressionLoggerImpl: async; chỉ log các item của trang trả về.",
            "Không có cột provenance / feature snapshot trên impression Social "
            "(khác Commerce Home).",
            "Không có bảng post_engage_event riêng; train dùng like/save/comment hiện hữu.",
        ],
    )

    add_heading(doc, "9. Pipeline offline (recsys-offline) — feed", 1)
    add_para(
        doc,
        "FastAPI app/main.py không phục vụ predict online. Các job feed (không gồm home-*):",
    )
    add_table(
        doc,
        ["Job", "Endpoint", "Vai trò"],
        [
            ["Clean", "POST /jobs/clean", "Extract Social Mongo/PG → CSV"],
            [
                "Build dataset",
                "POST /jobs/build-dataset",
                "Impression + PIT features + nhãn 24h → dataset.parquet",
            ],
            [
                "Split",
                "POST /jobs/split-dataset",
                "Time-ordered 80/10/10 theo shown_at",
            ],
            [
                "Train",
                "POST /jobs/train",
                "LightGBM binary → model.txt + train_meta.json",
            ],
            [
                "Evaluate",
                "POST /jobs/evaluate",
                "AUC + P/R/Hit@10 vs rule baseline → evaluate_report.json",
            ],
            [
                "Export-activate",
                "POST /jobs/export-activate",
                "ONNX + parity ≤1e-4 + insert model_artifacts + gate",
            ],
            [
                "Purchase profile",
                "POST /jobs/export-purchase-profile",
                "Commerce → user_purchase_profile.csv (cho feature offline)",
            ],
        ],
    )
    add_para(doc, "Nhãn và gate:", bold=True)
    add_bullets(
        doc,
        [
            "Positive nếu like hoặc save hoặc comment trong 24 giờ sau shown_at; không synthetic negative.",
            "Train: objective binary; learning_rate 0.05; num_leaves 31; boost 200; early stop 30.",
            "Activate nếu AUC LightGBM ≥ baseline và Precision@10 ≥ baseline; soft-reject giữ active cũ "
            "(exported_not_activated). Artifact: feed_ranker_v{N}.onnx.",
        ],
    )
    add_para(doc, "Simulation seed:", bold=True)
    add_bullets(
        doc,
        [
            "CLI: python -m simulation.cli (seed-db / simulate-memory / repair-post-timestamps).",
            "Bật bằng RECSYS_SIM_ALLOW=1; fashion-only; ghi impression/engage vào DB để train.",
            "Không gọi RecommendPosts online; không phải FastAPI /jobs/sim.",
        ],
    )

    add_heading(doc, "10. Frontend và Admin", 1)
    add_para(doc, "10.1 Web", bold=True)
    add_bullets(
        doc,
        [
            "Tab FEED_TABS.FOR_YOU = \"for-you\", nhãn \"Đề xuất\"; mặc định trên SocialFeedPage.",
            "fetchForYouFeed → GET /api/v1/social/feed/for-you; page size 20.",
        ],
    )
    add_para(doc, "10.2 Mobile (hiện trạng)", bold=True)
    add_bullets(
        doc,
        [
            "Tab nhãn “Đề xuất” nhưng id global → GET /feed/global (chronological).",
            "Chưa có fetchForYouFeed / wire /for-you.",
        ],
    )
    add_para(doc, "10.3 Model Registry (Admin)", bold=True)
    add_bullets(
        doc,
        [
            "GET /api/v1/social/admin/recommendation-model-artifacts?modelName= (mặc định feed_ranker)",
            "GET /api/v1/social/admin/recommendation-model-status",
            "Use case: ViewRecommendationModelArtifactsUseCase, ViewRecommendationModelStatusUseCase",
            "UI System Operations → Model registry: chọn Social For You (feed_ranker) "
            "vs Commerce Home (commerce_home_ranker); chỉ đọc, không activate từ UI.",
        ],
    )

    add_heading(doc, "11. Hằng số vận hành (đã khóa trong code)", 1)
    add_table(
        doc,
        ["Hằng số", "Giá trị"],
        [
            ["Feature dim", "6"],
            ["Candidate cap", "500"],
            ["Followee query cap", "300"],
            ["Recall windows", "7, 30, 90 ngày"],
            ["Min pool trước khi nới cửa sổ", "20"],
            ["Page size max", "50"],
            ["Recency half-life", "7 ngày"],
            ["Label window (offline)", "24 giờ"],
            ["Train objective", "binary LightGBM"],
            ["ONNX verify max abs diff", "1e-4"],
            ["Diversity Social", "Không"],
            ["Model name", "feed_ranker"],
        ],
    )

    add_heading(doc, "12. So sánh nhanh với Commerce Home", 1)
    add_table(
        doc,
        ["", "Social For You", "Commerce Home"],
        [
            ["Model", "feed_ranker", "commerce_home_ranker"],
            ["Features", "PostFeatureBuilder — 6", "HomeFeatureBuilder — 15"],
            ["Endpoint", "/social/.../feed/for-you", "/commerce/.../home/recommendations"],
            ["Diversity", "Không", "Có (greedy hard-cap)"],
            ["MODEL_ROOT", "SOCIAL_RECOMMENDATION_MODEL_ROOT", "COMMERCE_HOME_MODEL_ROOT"],
            ["Owner service", "social-service", "commerce-service"],
        ],
    )

    add_heading(doc, "13. Kết luận", 1)
    add_para(
        doc,
        "Chức năng đề xuất Social For You đã được triển khai end-to-end: recall ứng viên "
        "theo followee + public với fallback cửa sổ thời gian, vector 6 chiều, chấm điểm "
        "LightGBM ONNX (feed_ranker) hoặc rule-based khóa Phase-1, phân trang, ghi impression/"
        "seen, và pipeline offline train/evaluate/export-activate. Web đã gắn tab Đề xuất "
        "vào /for-you. Các hạn chế hiện hữu cần ghi nhận khi vận hành: cross-domain online "
        "đang stub rỗng, không diversity, mobile chưa gọi for-you, registry chỉ đọc.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
